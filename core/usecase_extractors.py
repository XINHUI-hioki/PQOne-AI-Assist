from docx import Document
from typing import List, Dict
from collections import Counter, defaultdict
from transformers import pipeline
import torch, os, logging, pytesseract, traceback, docx2txt
from statistics import mean
from datetime import datetime
import re, zipfile, html
from pathlib import Path
import pandas as pd
from PIL import Image
import re, numpy as np
print("[DEBUG] usecase_extractors loaded from", __file__)

_NOMINAL_V = 200.0 


os.environ["TOKENIZERS_PARALLELISM"] = "false" 
try:
    _blip = pipeline(
        "image-to-text",
        model="Salesforce/blip-image-captioning-base",   
        device=0 if torch.cuda.is_available() else -1
    )
except Exception as e:
    logging.warning(f"BLIP model failed to load: {e}")
    _blip = None 
def _read_docx_text(path: str | Path) -> str:
    """把 docx 的 word/document.xml 中所有 <w:t> 文本拼成纯文本"""
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    return "\n".join(
        html.unescape(t) for t in re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml)
    )

def translate_summary(summary: str, target_lang: str) -> str:
    """
    Translate summary string into the selected target language.

    Args:
        summary: The original English summary text
        target_lang: Expected values: "English", "中文", "日本語"

    Returns:
        Translated string (or original if target_lang is English)
    """
    if target_lang == "日本語":
        return "[日本語訳（仮）] " + summary  
    elif target_lang == "中文":
        return "[中文翻译（测试）] " + summary
    else:
        return summary 

def extract_worst_voltage_dip(docx_path: str) -> str:

    doc = Document(docx_path)
    worst_voltage = float('inf')
    worst_event = None

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]

            if len(cells) >= 8 and cells[2] == "Dip" and cells[3] == "IN":
                try:
                    time_str = cells[1]
                    channel = cells[4]
                    duration = cells[6]
                    voltage_str = cells[7]  # Worst

                    voltage = float(voltage_str.replace("V", ""))

                    if voltage < worst_voltage:
                        worst_voltage = voltage
                        worst_event = {
                            "time": time_str,
                            "channel": channel,
                            "duration": duration,
                            "voltage": voltage
                        }

                except Exception as e:
                    print(f"[❗Warning] Failed to parse row: {cells} ({e})")

    if worst_event:
        return (
            f"The worst voltage dip occurred at {worst_event['time']} "
            f"on Channel {worst_event['channel']}, lasted for {worst_event['duration']}, "
            f"and dropped to {worst_event['voltage']}V."
        )
    else:
        return "No voltage dip events were found in the document."

def extract_all_voltage_dips(docx_path: str | Path) -> pd.DataFrame:
    """
    Return a DataFrame with every “Dip IN” row in the report
    ── columns: time | channel | duration_ms | worst_voltage | depth_percent
    """
    doc = Document(docx_path)

    text = "\n".join(p.text for p in doc.paragraphs)
    m = re.search(r"Udin[^0-9]+([\d.]+)", text, flags=re.I)
    u_nom = float(m.group(1)) if m else _NOMINAL_V

    events: List[Dict] = []

    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 8 and cells[2] == "Dip" and cells[3] == "IN":
                try:
                    dur_val = float(re.sub(r"[^\d.]", "", cells[6]))  
                    worst_v = float(cells[7].replace("V", "").strip())
                    events.append({
                        "time": cells[1],
                        "channel": cells[4],
                        "duration_ms": dur_val,
                        "worst_voltage": worst_v,
                        "depth_percent": abs((u_nom - worst_v) / u_nom * 100),
                    })
                except Exception as e:
                    print(f"[⚠️] Skip broken Dip row: {cells} ({e})")

    if not events:
        return _extract_dips_by_regex(docx_path)

    return pd.DataFrame(events)

def _extract_dips_by_regex(docx_path: str | Path) -> pd.DataFrame:
    txt = _read_docx_text(docx_path)
    m = re.search(r"Udin[^0-9]+([\d.]+)", txt, flags=re.I)
    u_nom = float(m.group(1)) if m else _NOMINAL_V
    events: List[Dict] = []

    patt_tbl = re.compile(
        r"(?P<time>\d{2}:\d{2}:\d{2}\.\d{3})\s+Dip\s+IN\s+CH(?P<channel>\d)\s+"
        r"(?P<level>[\d.]+)V\s+(?P<duration>[\d.]+)\s*ms\s+(?P<worst>[-\d.]+)V",
        flags=re.I,
    )
    for m in patt_tbl.finditer(txt):
        worst_v = float(m.group("worst"))
        events.append({
            "time": m.group("time"),
            "channel": f"CH{m.group('channel')}",
            "duration_ms": float(m.group("duration")),
            "worst_voltage": worst_v,
            "depth_percent": abs((u_nom - worst_v) / u_nom * 100),
        })

    patt_short = re.compile(
        r"Duration:\s*(\d+)ms\s*/\s*Worst:\s*([-\d.]+)V",
        flags=re.I,
    )
    for dur, worst in patt_short.findall(txt):
        worst_v = float(worst)
        events.append({
            "time": None,
            "channel": None,
            "duration_ms": float(dur),
            "worst_voltage": worst_v,
            "depth_percent": abs((u_nom - worst_v) / u_nom * 100),
        })

    return pd.DataFrame(events)

def extract_all_voltage_dip_dicts(docx_path: str) -> List[Dict]:
    doc = Document(docx_path)
    dip_events = []

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 8 and cells[2] == "Dip" and cells[3] == "IN":
                try:
                    dip_event = {
                        "time": cells[1],
                        "channel": cells[4],
                        "duration": cells[6],
                        "voltage": float(cells[7].replace("V", ""))
                    }
                    dip_events.append(dip_event)
                except Exception as e:
                    print(f"[Warning] Failed to parse row: {cells} ({e})")

    return dip_events

def extract_worst_voltage_dip_dict(docx_path: str) -> Dict:
    """
    Extracts the worst voltage dip event from a Word report and returns a dict.
    """
    doc = Document(docx_path)
    worst_voltage = float('inf')
    worst_event = None

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 8 and cells[2] == "Dip" and cells[3] == "IN":
                try:
                    voltage = float(cells[7].replace("V", ""))
                    if voltage < worst_voltage:
                        worst_voltage = voltage
                        worst_event = {
                            "time": cells[1],
                            "channel": cells[4],
                            "duration": cells[6],
                            "voltage": voltage
                        }
                except Exception as e:
                    print(f"[⚠️] Failed to parse Dip row: {cells} ({e})")

    return worst_event if worst_event else {}

def extract_all_transients(docx_path: str) -> str:
    """
    Extracts and summarizes all Transient IN events from a Word report.
    Returns a natural-language summary string.
    """
    doc = Document(docx_path)
    events = []

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 8 and cells[2] == "Tran" and cells[3] == "IN":
                try:
                    events.append({
                        "time": cells[1],
                        "channel": cells[4],
                        "duration": cells[6],
                        "level": cells[5],
                        "worst": float(cells[7].replace("V", ""))
                    })
                except Exception as e:
                    print(f"[⚠️] Failed to parse Transient row: {cells} ({e})")

    if not events:
        return "No transient events were found in the document."

    description = "The following transient events were detected:\n"
    for e in events:
        description += (
            f"- At {e['time']} on Channel {e['channel']}, "
            f"a transient of level {e['level']} occurred, lasting {e['duration']}, "
            f"with a worst voltage of {e['worst']}V.\n"
        )
    return description.strip()

def extract_all_transient_dicts(docx_path: str) -> List[Dict]:
    doc = Document(docx_path)
    events = []

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 8 and cells[2] == "Tran" and cells[3] == "IN":
                try:
                    events.append({
                        "time": cells[1],
                        "channel": cells[4],
                        "duration": cells[6],
                        "level": cells[5],
                        "worst": float(cells[7].replace("V", ""))
                    })
                except Exception as e:
                    print(f"[⚠️] Failed to parse Transient row: {cells} ({e})")
    return events

def extract_all_inrush_events(docx_path: str) -> str:
    doc = Document(docx_path)
    events = []

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 8 and cells[2] == "Inrush" and cells[3] == "IN":
                try:
                    dur_str = cells[6] or "0"
                    dur_val = float(re.sub(r"[^\d.]", "", dur_str))  # "9" 或 "9 ms"
                    worst_val = float(cells[7].replace("A", "")) if cells[7] else None

                    events.append({
                        "time": cells[1],
                        "channel": cells[4],
                        "duration_ms": dur_val,
                        "level": cells[5],
                        "worst": worst_val,
                    })
                except Exception as e:
                    print(f"[⚠️] Skip broken Inrush row: {cells} ({e})")

    if not events:
        return "No inrush current events were found in the document."

    description = "The following inrush current events were detected:\n"
    for e in events:
        description += (
            f"- At {e['time']} on Channel {e['channel']}, "
            f"inrush level {e['level']} was observed, "
            f"lasting {e['duration_ms']} ms"
        )
        if e['worst'] is not None:
            description += f", with a worst current of {e['worst']} A"
        description += ".\n"

    return description.strip()

def extract_all_inrush_dicts(docx_path: str) -> List[Dict]:
    doc = Document(docx_path)
    events = []

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]

            if len(cells) >= 8 and cells[2] == "Inrush" and cells[3] == "IN":
                try:
                    duration_str = cells[6] or "0"
                    duration_ms = float(re.sub(r"[^\d.]", "", duration_str))

                    worst_str = cells[7].replace("A", "").strip()
                    worst_value = float(worst_str) if worst_str and worst_str != "---" else None

                    events.append({
                        "time": cells[1],
                        "channel": cells[4],
                        "level": cells[5],
                        "duration": f"{duration_ms} ms",
                        "worst": worst_value
                    })

                except Exception as e:
                    print(f"[⚠️] Failed to parse Inrush row: {cells} ({e})")

    return events

def extract_context_paragraphs(docx_path: str, min_length: int = 40, top_k: int = 3) -> str:
    """
    Extract top-K meaningful paragraphs from the Word report (non-table).
    Returns as readable context for summary or open question prompting.
    """
    try:
        min_length = int(min_length)
    except Exception:
        min_length = 40
    
    try:
        top_k = int(top_k)
    except Exception:
        top_k = 3
    doc = Document(docx_path)
    paragraphs = [
        para.text.strip()
        for para in doc.paragraphs
        if len(para.text.strip()) >= min_length
    ]

    selected = paragraphs[:top_k]
    if not selected:
        return "No meaningful paragraph content was found in the report."

    result = "The following are key paragraph excerpts from the report:\n\n"
    for i, para in enumerate(selected, 1):
        result += f"{i}. {para}\n\n"

    return result.strip()

def extract_voltage_dip_trend(doc: Document) -> str:
    """
    Analyzes all voltage dip events from the report and summarizes their time and channel distribution.
    Returns a natural-language description for trend analysis.
    """
    dip_events = []

    time_pattern = r'(\d{2}:\d{2}:\d{2}\.\d{3})'
    channel_pattern = r'CH(\d+)'
    
    for para in doc.paragraphs:
        text = para.text
        if "Dip" in text and re.search(time_pattern, text) and re.search(channel_pattern, text):
            time_match = re.search(time_pattern, text)
            channel_match = re.search(channel_pattern, text)
            if time_match and channel_match:
                time_str = time_match.group(1)
                ch = channel_match.group(1)
                try:
                    time_obj = datetime.strptime(time_str, "%H:%M:%S.%f")
                    dip_events.append((time_obj, ch))
                except ValueError:
                    continue

    if not dip_events:
        return "No voltage dip events found for trend analysis."

    time_bins = defaultdict(int)
    channel_count = defaultdict(int)

    for t, ch in dip_events:
        hour = t.hour
        if 6 <= hour < 12:
            time_bins["Morning"] += 1
        elif 12 <= hour < 18:
            time_bins["Afternoon"] += 1
        elif 18 <= hour < 24:
            time_bins["Evening"] += 1
        else:
            time_bins["Night"] += 1
        channel_count[ch] += 1

    # Build result string
    result = "Voltage Dip Trend Analysis:\n\n"
    result += f"Total number of dips: {len(dip_events)}\n\n"

    result += "Time distribution:\n"
    for period, count in time_bins.items():
        result += f"- {period}: {count} events\n"

    result += "\nChannel distribution:\n"
    for ch, count in channel_count.items():
        result += f"- Channel {ch}: {count} events\n"

    result += "\nBased on the above, please describe any noticeable patterns or trends in voltage dip occurrences."

    return result

def summarize_report(doc_path: str, language: str = "English") -> str:
    doc = Document(doc_path)

    dips = extract_all_voltage_dip_dicts(doc_path)
    trans = extract_all_transients(doc_path)
    inrush = extract_all_inrush_dicts(doc_path)
    worst = extract_worst_voltage_dip_dict(doc_path)  
    thd = extract_thd_metrics(doc_path)  
    harmonics_df, _ = extract_all_harmonics_from_docx(doc_path)
    harmonic_summary = {}

    if not harmonics_df.empty:
        for phase in harmonics_df['phase'].unique():
            df_phase = harmonics_df[harmonics_df['phase'] == phase]
            harmonic_summary[phase] = {
                int(row['order']): {
                    "avg": float(row['avg']),
                    "max": float(row['max'])
                }
                for _, row in df_phase.iterrows()
            }


    dip_count = len(dips)
    avg_voltage = round(mean([
        float(str(d['voltage']).replace("V", "")) for d in dips if "voltage" in d and d["voltage"] not in ["---", ""]
    ]), 2) if dips else None

    channels = [d['channel'] for d in dips]
    top_channel = max(set(channels), key=channels.count) if channels else None

    time_bins = Counter()
    for d in dips:
        try:
            try:
                hour = datetime.strptime(d['time'], "%H:%M:%S.%f").hour
            except ValueError:
                hour = datetime.strptime(d['time'], "%H:%M:%S").hour

            if 6 <= hour < 12:
                time_bins["Morning"] += 1
            elif 12 <= hour < 18:
                time_bins["Afternoon"] += 1
            elif 18 <= hour < 24:
                time_bins["Evening"] += 1
            else:
                time_bins["Night"] += 1
        except Exception as e:
            print(f"[WARN] Failed to parse time '{d.get('time', '')}': {e}")
            continue

    try:
        max_inrush = max([
            float(i['worst']) for i in inrush if i.get("worst") not in [None, "---", ""]
        ])
        max_inrush_str = f"{max_inrush:.2f} A"
    except:
        max_inrush_str = "N/A"

    name = doc_path.split("/")[-1]
    summary = f"📄 Report: {name}\n\n"
    summary += f"- Total voltage dips: {dip_count}\n"
    if avg_voltage:
        summary += f"- Average worst voltage: {avg_voltage} V\n"
    if top_channel:
        summary += f"- Most affected channel: {top_channel}\n"
    if time_bins:
        summary += "- Time distribution of dips:\n"
        for period, count in time_bins.items():
            summary += f"  • {period}: {count} events\n"

    summary += f"- Number of transient events: {len(trans)}\n"
    summary += f"- Number of inrush events: {len(inrush)}\n"
    summary += f"- Maximum inrush current: {max_inrush_str}\n"
    summary += f"- THD: U12 = {thd.get('U12')}, U23 = {thd.get('U23')}, U31 = {thd.get('U31')}\n"

    if isinstance(worst, dict):
        summary += (
            f"\nThe worst voltage dip occurred at {worst.get('time')} on Channel {worst.get('channel')}, "
            f"lasted for {worst.get('duration')}, and dropped to {worst.get('worst_voltage')}.\n"
        )

    summary = translate_summary(summary, language)
    return summary.strip()

def compare_multiple_reports(doc_paths: List[str]) -> str:
    summaries = [summarize_report(path) for path in doc_paths]

    text = "Comparison of Multiple PQONE Reports:\n\n"
    for i, summary_text in enumerate(summaries):
        text += f"--- Report {i+1} ---\n{summary_text}\n\n"

    text += "Based on the above summaries, please analyze any trends or differences across the reports."

    return text

def extract_u12_harmonics_description(docx_path: str) -> str:
    raw_text = docx2txt.process(docx_path)
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    thd_avg, thd_max = None, None
    harmonics = {}

    i = 0
    while i < len(lines) - 2:
        if lines[i].isdigit() and lines[i+1].endswith("V") and lines[i+2].endswith("V"):
            try:
                order = int(lines[i])
                avg = float(lines[i+1].replace("V", ""))
                max_ = float(lines[i+2].replace("V", ""))
                harmonics[order] = {"AVG": avg, "MAX": max_}
                i += 3
                continue
            except:
                pass
        if "THD" in lines[i]:
            segment = " ".join(lines[i:i+3])
            matches = re.findall(r"(\d+\.\d+)%", segment)
            if len(matches) >= 2:
                thd_avg = float(matches[0])
                thd_max = float(matches[1])
        i += 1

    if not harmonics:
        return "No U12 harmonics data were found in the document."

    desc = ""
    if thd_avg is not None and thd_max is not None:
        desc += f"The total harmonic distortion (THD) of U12 is {thd_avg}% (AVG) and {thd_max}% (MAX).\n"
        if thd_avg > 5.0:
            desc += "⚠️ Warning: The average THD exceeds 5%, which may indicate power quality issues.\n"

    desc += "Here are the extracted U12 harmonic voltage values:\n"
    for order in sorted(harmonics):
        v = harmonics[order]
        desc += f"- {order}th harmonic: {v['AVG']}V (AVG), {v['MAX']}V (MAX)\n"

    return desc.strip()

def extract_i12_harmonics_description(docx_path: str) -> str:
    doc_text = docx2txt.process(docx_path)

    thd_match = re.search(r"I12.*?THD\s*[:：]?\s*([\d.]+%)", doc_text)
    thd_text = f"THD of I12 is {thd_match.group(1)}." if thd_match else "I12 THD not found."

    pattern = re.compile(r"\b(\d{1,2})\b\s+I12\s+([\d.]+V)\s+([\d.]+V)")
    harmonics = pattern.findall(doc_text)

    if not harmonics:
        return f"{thd_text}\nNo harmonic data for I12 was found."

    description = f"{thd_text}\nI12 Harmonic Levels (order: avg → max):\n"
    for order, avg, peak in harmonics:
        description += f" - Order {order}: {avg} → {peak}\n"

    return description.strip()

def extract_p12_harmonics_description(docx_path: str) -> str:
    doc_text = docx2txt.process(docx_path)

    thd_match = re.search(r"P12.*?THD\s*[:：]?\s*([\d.]+%)", doc_text)
    thd_text = f"THD of P12 is {thd_match.group(1)}." if thd_match else "P12 THD not found."

    pattern = re.compile(r"\b(\d{1,2})\b\s+P12\s+([\d.]+V|W)\s+([\d.]+V|W)")
    harmonics = pattern.findall(doc_text)

    if not harmonics:
        return f"{thd_text}\nNo harmonic data for P12 was found."

    description = f"{thd_text}\nP12 Harmonic Levels (order: avg → max):\n"
    for order, avg, peak in harmonics:
        description += f" - Order {order}: {avg} → {peak}\n"

    return description.strip()

def extract_all_rvc_events(doc_path):
    """从 Word 表格中提取所有 RVC 事件行"""
    doc = Document(doc_path)
    rvc_events = []

    for table in doc.tables:
        for row in table.rows[1:]:
            cells = row.cells
            try:
                event_type = cells[2].text.strip()
                if event_type == "RVC":
                    time = cells[1].text.strip()
                    direction = cells[3].text.strip()
                    channel = cells[4].text.strip()
                    level = cells[5].text.strip()
                    duration = cells[6].text.strip()
                    worst = cells[7].text.strip()

                    rvc_events.append({
                        "time": time,
                        "direction": direction,
                        "channel": channel,
                        "voltage": level if level else "N/A",
                        "duration": duration if duration else "N/A",
                        "delta_u": worst if worst else "N/A"
                    })
            except IndexError:
                continue  # skip broken rows

    return rvc_events

def extract_thd_metrics(doc_path: str) -> Dict[str, str]:
    """
    从 extract_all_harmonics_from_docx 中获取 THD 字典。
    返回格式：{channel_name: "x.xx%"}
    """
    _, thd_dict = extract_all_harmonics_from_docx(doc_path)

    if not thd_dict:
        print(f"[INFO] No THD values found in document: {doc_path}")
        return {}

    return thd_dict

def summarize_reports_as_table(doc_paths: List[str]) -> str:

    rows = []
    for path in doc_paths:
        name = path.split("/")[-1]
        dips = extract_all_voltage_dip_dicts(path)
        inrush = extract_all_inrush_dicts(path)
        thd = extract_thd_metrics(path)

        dip_count = len(dips)
        try:
            avg_voltage = round(mean([
                float(str(d['voltage']).replace("V", "")) for d in dips if "voltage" in d and d["voltage"] not in ["", "---"]
            ]), 2)
            avg_voltage_str = f"{avg_voltage} V"
        except:
            avg_voltage_str = "N/A"

        try:
            max_inrush = max([
                float(i["worst"]) for i in inrush if i.get("worst") not in ["", None, "---"]
            ])
            max_inrush_str = f"{max_inrush:.2f} A"
        except:
            max_inrush_str = "N/A"

        rows.append([
            name,
            str(dip_count),
            avg_voltage_str,
            max_inrush_str,
            thd.get("U12", "N/A"),
            thd.get("U23", "N/A"),
            thd.get("U31", "N/A")
        ])

    header = "| Report Name | Dips | Avg Worst V | Max Inrush | U12 THD | U23 THD | U31 THD |"
    sep =    "|-------------|------|--------------|-------------|---------|---------|---------|"
    table = [header, sep]
    for row in rows:
        table.append("| " + " | ".join(row) + " |")

    return " Power Quality Summary:\n\n" + "\n".join(table)

def get_harmonics_df(docx_path: str) -> pd.DataFrame:
    return extract_all_harmonics_from_docx(docx_path)

def get_dip_df(docx_path) -> pd.DataFrame:
    """Return cleaned dip table with 'time', 'duration_ms', 'worst_voltage'."""
    rows = extract_all_voltage_dip_dicts(docx_path)  
    if not rows:
        return pd.DataFrame()

    df = pd.DataFrame(rows)
    rename_map = {}

    for col in df.columns:
        c = col.lower().strip()

        if any(k in c for k in ["time", "date/time", "timestamp"]):
            rename_map[col] = "time"

        elif "duration" in c:
            rename_map[col] = "duration_ms"

        elif re.search(r"(worst|level|value|voltage)", c):
            rename_map[col] = "worst_voltage"

    df = df.rename(columns=rename_map)

    if "time" in df.columns:
        df["time"] = pd.to_datetime(df["time"], errors="coerce")

    for c in ["worst_voltage", "duration_ms"]:
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce")

    if "worst_voltage" not in df.columns:
        print(f"[WARN] get_dip_df: 'worst_voltage' column missing after renaming.")
        return pd.DataFrame()

    return df.dropna(subset=["worst_voltage"]).reset_index(drop=True)

def get_thd_trend_df(docx_path: str) -> pd.DataFrame:
    """
    Parse “[ Harmonics - Trend ]” 表，返回含
    ['time', 'thd'] 两列的 DataFrame，供折线图使用。
    """
    doc = Document(docx_path)
    trend_rows: list[tuple[str, float]] = []

    target_tbl = None
    for tbl in doc.tables:
        first_row = " ".join(cell.text for cell in tbl.rows[0].cells)
        if "Harmonics - Trend" in first_row:
            target_tbl = tbl
            break
    if target_tbl is None:
        print("[INFO] THD trend table not found — skip this chart.")
        return None            


    for row in target_tbl.rows[2:]:                     
        cells = [c.text.strip() for c in row.cells[:2]] 
        if len(cells) < 2 or not cells[0]:
            continue
        time_txt, thd_txt = cells
        try:
            tm = datetime.strptime(time_txt, "%H:%M:%S.%f")
        except ValueError:
            continue
        thd_val = float(re.sub(r"[^\d.]", "", thd_txt)) if thd_txt else None
        if thd_val is not None:
            trend_rows.append((tm.replace(year=1900, month=1, day=1), thd_val))

    if not trend_rows:
        raise ValueError("No THD trend rows could be parsed.")

    return pd.DataFrame(trend_rows, columns=["time", "thd"])

def get_thd_avg_df(docx_path: str) -> pd.DataFrame | None:
    _, thd_result = extract_all_harmonics_from_docx(docx_path)
    
    if not thd_result:
        print("[INFO] No THD data found in harmonics table.")
        return None

    rows = []
    for phase, vals in thd_result.items():
        if "avg" in vals:
            try:
                num_val = float(vals["avg"].replace("%", ""))
                rows.append({"phase": phase, "THD_avg": num_val})
            except ValueError:
                continue

    if not rows:
        print("[INFO] No valid THD-avg values found.")
        return None

    return pd.DataFrame(rows)

def get_thd_avgmax_df(path: str) -> pd.DataFrame | None:
    """
    输出列：phase | THD_avg | THD_max
    """
    _, thd_dict = extract_all_harmonics_from_docx(path)
    rows = []
    for phase, d in thd_dict.items():
        row = {"phase": phase}
        try:
            row["THD_avg"] = float(d.get("avg", "nan").rstrip("%"))
        except Exception:
            row["THD_avg"] = np.nan
        try:
            row["THD_max"] = float(d.get("max", "nan").rstrip("%"))
        except Exception:
            row["THD_max"] = np.nan
        rows.append(row)
    return pd.DataFrame(rows) if rows else None

def _parse_thd_row(cells, first_row, thd_result):
    """
    cells  : ['THD', '6.27%', '7.47%', '5.13%', '6.80%', ...]
    first_row: ['Order', 'U12', '', 'U23', '', ...]
               ↑        ↑ col1  ↑ col2   ↑ col3  ↑ col4
               │每两列为同一相
    """
    # 计算相别数
    pairs = (len(cells) - 1) // 2          # 几个 AVG+MAX 对
    for p in range(pairs):
        # col indexes
        col_avg = 1 + 2 * p
        col_max = col_avg + 1

        phase = first_row[col_avg].split("\n")[0].strip() or f"CH{p+1}"
        avg_val = cells[col_avg]
        max_val = cells[col_max]

        if phase not in thd_result:
            thd_result[phase] = {}
        if re.search(r"\d", avg_val):
            thd_result[phase]["avg"] = re.search(r"(\d+(\.\d+)?)", avg_val).group(1)
        if re.search(r"\d", max_val):
            thd_result[phase]["max"] = re.search(r"(\d+(\.\d+)?)", max_val).group(1)

def analyze_uploaded_image(image_path: str, lang: str = "English") -> str:
    print("[DEBUG] controller got image path =", image_path)
    print("[DEBUG] file exists?", os.path.exists(image_path))

    caption = ""
    if _blip:
        try:
            caption = _blip(image_path, max_new_tokens=40)[0]["generated_text"]
        except Exception as e:
            logging.warning(f"[OCR] Caption failed: {e}")

    try:
        ocr_text = pytesseract.image_to_string(Image.open(image_path))
    except Exception as e:
        print(" OCR failed with exception:")
        traceback.print_exc()  
        ocr_text = f" OCR failed: {e}"  

    summary = (
        f"**Image caption**: {caption}\n"
        f"**OCR text**:\n{ocr_text[:500]}..."
    )
    return translate_summary(summary, lang)

def analyze_file(doc_path: str, language: str = "English") -> str:
    """
    Quick one-page summary used by the 'analyze file' use-case.
    Leverages existing helpers so there’s almost no new parsing code.
    """
    worst   = extract_worst_voltage_dip(doc_path)
    thd     = extract_thd_metrics(doc_path)
    dip_cnt = len(extract_all_voltage_dip_dicts(doc_path))
    inrush  = extract_all_inrush_dicts(doc_path)
    rvc_cnt = len(extract_all_rvc_events(doc_path))
    df, thd = extract_all_harmonics_from_docx(doc_path)
    thd_lines = ""
    if thd:
        thd_lines = "- Harmonic THD:\n"
        for ch, vals in thd.items():
            avg_str = vals.get("avg", "N/A")
            max_str = vals.get("max", "N/A")
            thd_lines += f"  • {ch} — AVG: {avg_str}, MAX: {max_str}\n"
    else:
        thd_lines = "- Harmonic THD: Not found\n"
    harmonics_desc = ""
    if not df.empty:
        sampled = df[df["order"].isin([3, 5, 7])]  # 只挑几个次谐波
        grouped = sampled.groupby("phase")
        for phase, group in grouped:
            desc = ", ".join(
                f"{int(row['order'])}th={row['avg']:.2f}%" for _, row in group.iterrows()
            )
            harmonics_desc += f"_ Harmonics ({phase}): {desc}\n"
    else:
        harmonics_desc = "_ Harmonics: Not available\n"
    name = Path(doc_path).name
    summary = (
        f" **{name}**\n"
        f"- Voltage-dip events: **{dip_cnt}**\n"
        f"- Worst dip: {worst}\n"
        f"{thd_lines}"
        f"- Inrush events: {len(inrush)}\n"
        f"- RVC events: {rvc_cnt}\n"
        f"{harmonics_desc}"
    )
    return translate_summary(summary, language)

def extract_all_harmonics_from_docx(docx_path: str) -> pd.DataFrame:
    doc = Document(docx_path)
    records = []
    thd_result = {}

    for table in doc.tables:
        first_row = [cell.text.strip() for cell in table.rows[0].cells]
        second_row = [cell.text.strip() for cell in table.rows[1].cells] if len(table.rows) > 1 else []

        if not first_row or not any(re.search(r'\d', c) for c in first_row):
            continue  

        for row in table.rows[2:]:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) < 3:
                continue

            first_cell = cells[0].strip().lower().replace(" ", "").replace("\n", "")

            if "thd" in first_cell:
                _parse_thd_row(cells, first_row, thd_result)
                continue          # THD 行处理完就跳过，进入下一行

            if not re.match(r'^\d+$', cells[0]):
                continue  

            try:
                order = int(cells[0])
                for i in range(1, len(cells) - 1, 2):
                    ch_name = first_row[i] if i < len(first_row) else f"CH{i}"
                    avg = float(re.sub(r"[^\d.]+", "", cells[i]))
                    max_ = float(re.sub(r"[^\d.]+", "", cells[i + 1]))
                    records.append({
                        "phase": ch_name,
                        "order": order,
                        "avg": avg,
                        "max": max_
                    })
            except Exception as e:
                print(f"[WARN] Skip the abnormal line {cells}: {e}")
                continue
    
    if not records:
        print("[INFO] No harmonic data was extracted")
        return pd.DataFrame()

    return pd.DataFrame(records), thd_result


def _check_dip_alert(dips: list[dict], threshold: float) -> str | None:
    alert_events = [d for d in dips if d.get("voltage") and float(d["voltage"]) < threshold]
    if alert_events:
        worst = min(alert_events, key=lambda d: float(d["voltage"]))
        return (f"⚠️ **Dip Alert** – {len(alert_events)} events below {threshold} V.  "
                f"Worst-case {worst['voltage']} V at {worst['time']} ({worst['channel']}).")
    return None

def _check_thd_alert(thd_dict: dict, threshold: float) -> str | None:
    over = {ph: vals for ph, vals in thd_dict.items()
            if vals and float(str(vals.get('avg','0')).rstrip('%')) > threshold}
    if over:
        lines = ", ".join(f"{ph} {vals['avg']}" for ph, vals in over.items())
        return f"⚠️ **THD Alert** – Avg THD over {threshold}% on {lines}."
    return None


def _extract_metrics_for_compare(doc_path: str) -> dict:
    """抽几项最关心的指标返回字典，供比较用。"""
    worst_dip  = extract_worst_voltage_dip_dict(doc_path)
    thd        = extract_thd_metrics(doc_path)
    inrush     = extract_all_inrush_dicts(doc_path)
    name       = doc_path.split("/")[-1]

    return {
        "name": name,
        "worst_dip_v": worst_dip.get("worst_voltage") if worst_dip else "N/A",
        "worst_dip_ch": worst_dip.get("channel") if worst_dip else "N/A",
        "dip_cnt": len(extract_all_voltage_dip_dicts(doc_path)),
        "trans_cnt": len(extract_all_transients(doc_path)),
        "inrush_max": max(
            [float(i.get("worst", 0)) for i in inrush if i.get("worst") not in [None, "", "---"]],
            default=0
        ),
        "thd_u12": thd.get("U12"),
        "thd_u23": thd.get("U23"),
        "thd_u31": thd.get("U31"),
    }

    