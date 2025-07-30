from __future__ import annotations

import json
import queue
import threading
from datetime import datetime
from docx import Document        
from statistics import mean
from typing import List, Sequence
from rag.RAG_background import qa_chain
import sys
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from globals import GLOBAL
from globals import lc_memory
from model_io import call_language_model_stream, rag_search
from rag.RAG_background import (
    llm,
    generate_followup_question,
    conversation_chain,
    memory,
)
from usecase_extractors import (
    extract_context_paragraphs,
    extract_worst_voltage_dip,
    extract_all_voltage_dips,
    extract_all_transients,
    extract_all_inrush_events,
    extract_voltage_dip_trend,
    extract_all_voltage_dip_dicts,
    extract_all_transient_dicts,
    extract_all_inrush_dicts,
    extract_worst_voltage_dip_dict,
    extract_u12_harmonics_description,
    extract_i12_harmonics_description,
    extract_p12_harmonics_description,
    extract_all_rvc_events,
    extract_thd_metrics,
    summarize_report,
    analyze_file, 
    analyze_uploaded_image, 
    get_harmonics_df ,
    _check_dip_alert ,
    _check_thd_alert ,
    extract_all_harmonics_from_docx,
    _extract_metrics_for_compare
)

'''
Receive "User Questions", "file paths", "parameters", etc. sent by app_gui
The internal start_generation_thread() starts the background thread and calls _generate_in_thread()
_generate_in_thread() is responsible for:
Problem preprocessing, use case identification (such as dip, summary, harmonic, etc.)
Call **usecase_extractors.py** to perform event/structured extraction
If you need to retrieve supplementary materials, call rag_search() of model_io.py.
Concatenate the prompt and call call_language_model_stream() of [model_io.py] to complete the final question answering
Support followup intelligent question generation, conversation memory, etc. as needed

'''

import datetime

_BEGINNER_KEYWORDS = [
    "open", "load", "read", "how to", "where is", "view",
    "trend graph", "data list", "event", "report", "toolbar", "screen", "drag", "drop"
]

def _is_beginner_question(text: str) -> bool:
    lower = text.lower()
    return any(keyword in lower for keyword in _BEGINNER_KEYWORDS)


def build_prompt(extracted_info: str, background_context: str, user_question: str, lang: str):
    lang_instruction = {
        "English": "Please answer in English.",
        "中文": "请用简体中文回答。",
        "日本語": "日本語で答えてください。",
    }.get(lang, "Please answer in English.")

    instruction_map = {
        "English": (
            "You are HiokiAssist, an intelligent assistant for power quality analysis.\n"
            "Your task is to answer the user's question **based solely on the context provided below**.\n"
            "Use reasoning, summarize the relevant facts, and be concise.\n"
            "If the answer is in a table, you may quote or summarize it.\n"
            "If you cannot find the answer, say \"I can't find it in the file.\"\n\n"
            "🚫 IMPORTANT: Do NOT comment on, discuss, or compare any non-Hioki brands such as Fluke, Keysight, NI, Omron, etc.\n"
            "If the user's question is about other companies, politely reply:\n"
            "\"I'm sorry, I can only assist with Hioki products and services.\""
        ),
        "中文": (
            "你是 HiokiAssist，一名用于电能质量分析的智能助手。\n"
            "你的任务是**仅根据下方提供的上下文**回答用户的问题。\n"
            "请进行推理、总结关键事实，并保持简洁清晰。\n"
            "如答案在表格中，可引用或概括。\n"
            "如果找不到答案，请回复“文件中没有相关信息”。\n\n"
            "🚫 重要提示：请勿评论、讨论或比较任何非 Hioki 品牌，例如 Fluke、Keysight、NI、Omron 等。\n"
            "如果用户的问题涉及其他公司，请礼貌地回复：\n"
            "“很抱歉，我只能协助解答与 Hioki 产品和服务相关的问题。”"
        ),
        "日本語": (
            "あなたは HiokiAssist、電力品質解析用のインテリジェントアシスタントです。\n"
            "以下のコンテキストに**基づいてのみ**、ユーザーの質問に答えてください。\n"
            "推論を行い、関連する事実を要約し、簡潔に答えてください。\n"
            "回答が表にある場合は、引用または要約しても構いません。\n"
            "情報が見つからない場合は、「ファイルに関連情報が見つかりませんでした」と答えてください。\n\n"
            "🚫 注意：Fluke、Keysight、NI、Omron などの Hioki 以外のブランドについてコメント・議論・比較しないでください。\n"
            "他社に関する質問があった場合は、次のように丁寧に返答してください：\n"
            "「申し訳ありませんが、Hioki 製品およびサービスに関するご質問のみお手伝いできます。」"
        ),
    }

    instruction = instruction_map.get(lang, instruction_map["English"])
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    context_parts = []
    if extracted_info:
        context_parts.append("Information extracted from uploaded report:\n" + extracted_info)
    if background_context:
        context_parts.append("Related background knowledge:\n" + background_context)
    full_context = "\n\n".join(context_parts) if context_parts else "No relevant context found."

    # 最终 prompt 拼接
    prompt = (
        f"{lang_instruction}\n\n"
        f"{instruction}\n"
        f"Current time: {timestamp}\n\n"
        "Context:\n"
        f"{full_context}\n\n"
        "User question:\n"
        f"{user_question.strip()}\n\n"
        f"{lang_instruction}\n\n"
        "At the end of your answer, suggest one helpful next question."
    )
    return prompt


def preprocess_query(query: str) -> str:
    query = query.lower().strip()
    if query.endswith("?"):
        query = query[:-1]
    if query.startswith("tell me"):
        query = query.replace("tell me", "what is", 1)
    elif query.startswith("can you tell me"):
        query = query.replace("can you tell me", "what is", 1)
    elif query.startswith("could you explain"):
        query = query.replace("could you explain", "explain", 1)
    return query

def detect_usecase_from_question(question: str) -> str:
    q = question.lower()
    if any(k in q for k in ["summary", "summarize", "overall"]):
        return "summary"
    if any(k in q for k in ["worst voltage", "lowest voltage", "worst dip"]):
        return "worst_voltage"
    if any(k in q for k in ["all dips", "list dips", "dip events","voltage dips", "voltage dip", "dip events",
    "dip duration", "dip table", "list dips", "dip summary"]):
        return "dips"
    if "transient" in q or "spike" in q:
        return "transients"
    if "inrush" in q or "startup current" in q:
        return "inrush"
    if any(k in q for k in ["trend", "pattern", "distribution"]):
        return "trend_dips"
    if "harmonic" in q  or "distortion" in q:
        if "current" in q or "i12" in q:
            return "current_harmonic"
        if "power" in q or "p12" in q:
            return "power_harmonic"
        return "harmonic"
    if "rvc" in q or "rapid voltage" in q:
        return "rvc_events"
    if "thd" in q or "harmonic distortion" in q:
        return "thd_metrics"
    if ("analyze" in q and "image" in q) or ("analyze" in q and "picture" in q):
        return "analyze_image"

    if any(k in q for k in ["compare", "difference", "which report"]):
        return "compare"
    return "paragraphs"

def summarize_dips(dips):
    if not dips:
        return "No voltage dips were detected in the report."
    avg_voltage = round(mean(d["voltage"] for d in dips), 2)
    count = len(dips)
    top_channel = max({d["channel"] for d in dips}, key=[d["channel"] for d in dips].count)
    return (
        f"A total of {count} voltage dips were recorded.\n"
        f"The average worst voltage across all dips was {avg_voltage} V.\n"
        f"The most frequently affected channel was Channel {top_channel}."
    )

def summarize_transients(trans):
    if not trans:
        return "No transient events were found in the report."
    first = trans[0]
    return (
        f"A total of {len(trans)} transient events were detected.\n"
        f"The first transient occurred at {first['time']} on Channel {first['channel']}."
    )

def summarize_inrush(inrush):
    if not inrush:
        return "No inrush current events were found in the report."
    max_current = max(i["worst"] for i in inrush)
    return (
        f"{len(inrush)} inrush current events were recorded in total.\n"
        f"The highest recorded inrush current was {max_current:.2f} A."
    )

def summarize_worst_dip(worst):
    if not worst:
        return "No worst voltage dip was found in the report."
    return (
        f"The worst voltage dip occurred at {worst['time']} on Channel {worst['channel']}, "
        f"lasting {worst['duration']} and dropping to {worst['voltage']} V."
    )

def summarize_rvc(events):
    from collections import Counter
    if not events:
        return "No RVC events found in the report."
    lines = [f"- {e['time']} | {e['direction']} | {e['channel']} | ΔU: {e['delta_u']}" for e in events[:5]]
    counts = Counter(e["channel"] for e in events)
    by_channel = "\n".join(f"- {ch}: {cnt} events" for ch, cnt in counts.items())
    try:
        max_delta = max(events, key=lambda e: abs(float(e["delta_u"].rstrip("V"))))
        extreme = f"\nMax ΔU: {max_delta['delta_u']} at {max_delta['time']} on {max_delta['channel']}"
    except Exception:  
        extreme = ""
    return "RVC Summary:\n" + "\n".join(lines) + "\n\nRVC Event Count by Channel:\n" + by_channel + extreme


def _generate_in_thread(
    prompt: str,
    lang: str,
    model_name: str,
    memory_mode: bool,
    doc_paths: List[str],
    df_dict: dict | None,
    ui_queue: queue.Queue,
    params: dict,
) -> None:
    try:
        def get_cached(key: str, fn, *args):

            if df_dict is None:
                return fn(*args)                
            cache = df_dict.setdefault(primary_doc, {})
            if key not in cache:
                cache[key] = fn(*args)
            return cache[key]        
        user_question = preprocess_query(prompt)
        dip_th  = float(params.get("dip_th", 180))
        thd_th  = float(params.get("thd_th", 5))

        if GLOBAL.config.beginner_mode:
            from rag.RAG_background import search_md_only_context
            relevant_docs = search_md_only_context(prompt)
            context = "\n\n".join([doc.page_content for doc in relevant_docs])

            print("\n===== BEGINNER MODE CONTEXT =====\n")
            print(context, "...\n")         

            system_prompt_map = {
                "English": (
                    "Answer the user's question strictly in English using the provided documentation.If there are any pictures, please show them\n\n"
                    "You are a helpful assistant for explaining the PQ ONE software.\n"
                    "You are a Hioki-specific assistant. Do not mention or comment on products from other companies.\n"
                    "If the question is not about Hioki or electrical theory, say you cannot help.\n\n"
                    "Use the following documentation snippets to answer the user's question thoroughly and completely.\n"
                    "If the documentation describes multiple items (such as 5 areas), list and explain each one clearly.If have image, please show\n"
                    "If the answer is not found in the documents, say: \"I can't find it in the file.\"\n"
                ),
                "中文": (
                    "请严格依据以下文档内容，用简体中文回答用户问题。如果有图片，请你展示图片。\n\n"
                    "你是 Hioki PQ ONE 软件的助手。\n"
                    "你只回答与 Hioki 相关的问题，禁止评论、比较其他品牌（如 Fluke、Keysight、NI 等）。\n"
                    "如果问题与 Hioki 或电学无关，请回复：“我无法回答这个问题”。\n\n"
                    "请详细解释文档中涉及的每一点内容（如有 5 个区域则分别列出）。\n"
                    "如果文件中没有提到相关内容，请回复：“文件中没有相关信息”。\n"
                ),
                "日本語": (
                    "以下のドキュメントに基づき、日本語でユーザーの質問に答えてください。写真があれば見せてください\n\n"
                    "あなたは Hioki PQ ONE ソフトウェアのアシスタントです。\n"
                    "Hioki に関係する質問のみに答えてください。他社（Fluke、Keysight、NI など）については回答しないでください。\n"
                    "Hioki または電気に関係のない質問には「その質問にはお答えできません」と答えてください。\n\n"
                    "複数の項目がある場合（例：5つのエリア）には、それぞれを明確に説明してください。\n"
                    "文書に情報がない場合は「ファイルに関連情報が見つかりませんでした」と答えてください。\n"
                )
            }
            system_prompt = system_prompt_map.get(lang, system_prompt_map["English"]) + context


            print("\n===== BEGINNER MODE PROMPT =====\n")
            print(system_prompt, "...\n")   
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt}
            ]

            answer = call_language_model_sync(
                messages=messages,
                temperature=0.4,
                model_name="llama3",     
                max_tokens=500
            )

            ui_queue.put(answer)
            ui_queue.put(None)
            return

        usecase = detect_usecase_from_question(user_question)

        history_block = ""
        if memory_mode:
            msgs = lc_memory.chat_memory.messages[-6:]     
            history_block = "\n\n".join(
                f"{m.type.capitalize()}: {m.content}" for m in msgs
            )


        extracted_info = ""
        if doc_paths:
            primary_doc = doc_paths[0]

            if usecase == "summary":
                dips = []          
                trans = []
                inrush = []
                rvc = []
                thd = {}

                dips = get_cached("dips", extract_all_voltage_dip_dicts, primary_doc)
                trans = get_cached("trans", extract_all_transient_dicts, primary_doc)
                inrush = get_cached("inrush", extract_all_inrush_dicts, primary_doc)
                worst = get_cached("worst_dip", extract_worst_voltage_dip_dict, primary_doc)
                harm_df, _ = get_cached("u_h", extract_all_harmonics_from_docx, primary_doc)
                u_h = "Harmonics data extracted ({} rows)".format(len(harm_df))
                i_h = extract_i12_harmonics_description(primary_doc)
                p_h = extract_p12_harmonics_description(primary_doc)
                rvc = get_cached("rvc", extract_all_rvc_events, primary_doc) 
                thd = get_cached("thd", extract_thd_metrics, primary_doc) 
                warnings = []
                if dips:                     
                    dip_warn = _check_dip_alert(dips, dip_th)
                    if dip_warn:
                        warnings.append(dip_warn)

                if thd:                   
                    thd_warn = _check_thd_alert(thd, thd_th)
                    if thd_warn:
                        warnings.append(thd_warn)
                extracted_info = "\n\n".join([
                    *warnings, 
                    summarize_dips(dips),
                    summarize_worst_dip(worst),
                    summarize_transients(trans),
                    summarize_inrush(inrush),
                    u_h, i_h, p_h,
                    summarize_rvc(rvc),
                    f"THD Summary:\n- U12 THD: {thd.get('U12')}\n- U23 THD: {thd.get('U23')}\n- U31 THD: {thd.get('U31')}",
                ])

            elif usecase == "worst_voltage":
                extracted_info = get_cached("worst_voltage",extract_worst_voltage_dip, primary_doc) 

            elif usecase == "dips":
                dips_df = dips_df = get_cached("dips_df", extract_all_voltage_dips, primary_doc)
                extracted_info = (
                    f"Total voltage-dip events: {len(dips_df)}\n\n"
                    + dips_df.to_markdown(index=False)
                )

            elif usecase == "transients":
                extracted_info = get_cached("transient", extract_all_transients, primary_doc) 

            elif usecase == "inrush":
                extracted_info = get_cached("inrush", extract_all_inrush_events, primary_doc) 

            elif usecase == "analyze_file":
                if df_dict and primary_doc in df_dict:
                    tables = df_dict[primary_doc]
                    thd = tables.get("thd_avg")
                    dips = tables.get("dip")
                    harmonics = tables.get("harmonics")
                    extracted_info = "Quick summary (cached tables):\n"

                    if dips is not None and not dips.empty:
                        extracted_info += f"- Dip events: {len(dips)}\n"
                        extracted_info += f"- Avg worst voltage: {dips['worst_voltage'].mean():.2f} V\n"
                    if thd is not None and not thd.empty:
                        extracted_info += (
                            f"- Avg THD trend points: {len(thd)}\n"
                            f"- Max THD value: {thd['thd'].max():.2f}%\n"
                        )
                    if harmonics is not None and not harmonics.empty:
                        extracted_info += f"- Harmonic orders captured: {harmonics['order'].nunique()}\n"
                else:
                    extracted_info = analyze_file(primary_doc, lang)

            elif usecase == "analyze_image":
                extracted_info = analyze_uploaded_image(primary_doc, lang)

            elif usecase == "trend_dips":
                if primary_doc.lower().endswith((".docx", ".doc")):
                    doc_obj = Document(primary_doc)
                    extracted_info = get_cached("trend_dips", extract_voltage_dip_trend(doc_obj)) 
                else:
                    extracted_info = "Trend extraction仅支持 .docx 报告文件。"

            elif usecase == "harmonic":
                df = get_cached("harmonics", get_harmonics_df, primary_doc) 
                if df is None or df.empty:
                    extracted_info = "No harmonic data was found in the report."
                else:
                    table = df[df["order"].isin(sorted(df["order"].unique()))]
                    extracted_info = f"Extracted harmonics (top orders):\n\n{table.to_markdown(index=False)}"

            elif usecase == "current_harmonic":
                extracted_info = get_cached("current_harmonics", extract_i12_harmonics_description, primary_doc)

            elif usecase == "power_harmonic":
                extracted_info = get_cached("power_harmonics", extract_p12_harmonics_description, primary_doc) 

            elif usecase == "rvc_events":
                rvc_events = get_cached("rvc_events", extract_all_rvc_events, primary_doc) 
                extracted_info = "\n".join(
                    f"{e['time']} | {e['direction']} | {e['channel']} | ΔU: {e['delta_u']}"
                    for e in rvc_events
                ) or "No RVC events found in the report."

            elif usecase == "thd_metrics":
                thd = get_cached("thd_metrics", extract_thd_metrics, primary_doc) 
                if not thd:
                    extracted_info = "No THD values were found in the report."
                else:
                    extracted_info = "THD Analysis:\n" + "\n".join(
                        f"- {ch} THD: {val}" for ch, val in thd.items()
                    )

            elif usecase == "compare":
                metrics_list = [_extract_metrics_for_compare(p) for p in doc_paths]

                header = "| Report | Worst-Dip (V) | Channel | #Dips | #Trans | Max Inrush (A) | THD U12 | THD U23 | THD U31 |\n" \
                        "|--------|--------------|---------|-------|--------|----------------|---------|---------|---------|\n"
                rows = []
                for m in metrics_list:
                    rows.append(
                        f"| {m['name']} | {m['worst_dip_v']} | {m['worst_dip_ch']} | "
                        f"{m['dip_cnt']} | {m['trans_cnt']} | {m['inrush_max']} | "
                        f"{m['thd_u12']} | {m['thd_u23']} | {m['thd_u31']} |"
                    )
                compare_table = header + "\n".join(rows)

                individual_summaries = "\n\n".join(summarize_report(p, lang) for p in doc_paths)

                extracted_info = (
                    "### 🔍 Key Metrics Comparison\n\n"
                    + compare_table
                    + "\n\n---\n\n"
                    + individual_summaries
                )

            else:
                extracted_info = get_cached(" compare", extract_context_paragraphs, primary_doc) 

        lower_q = user_question.lower()
        rag_needed = any(k in lower_q for k in ["how to", "operate", "step", "navigate", "open", "save", "export", "setting", "set"])
        background_context = ""
        rag_future = None

        if rag_needed or not extracted_info:
            from concurrent.futures import ThreadPoolExecutor
            executor = ThreadPoolExecutor(max_workers=1)
            rag_future = executor.submit(rag_search, user_question, lang)

        try:
            background_context = rag_future.result(timeout=12) or ""
        except Exception as e:
            print(f"[WARN] RAG search failed: {e}")
            background_context = ""
        if not extracted_info and not background_context:
            extracted_info = "I can't find relevant info in file or KB."

        context_parts = []
        if extracted_info:
            context_parts.append("Information extracted from uploaded report:\n" + extracted_info)
        if background_context:
            context_parts.append("Related background knowledge:\n" + background_context)
        if memory_mode and history_block:
            context_parts.append(history_block)
        if memory_mode and lc_memory.chat_memory.messages:
            previous_dialogue = format_lc_history(lc_memory)
            user_question = previous_dialogue + "\nUser: " + user_question


        full_prompt = build_prompt(extracted_info, background_context, user_question, lang)

        options = {
            "temperature": float(params.get("temp", 0.8)),
            "top_p": float(params.get("top_p", 0.9)),
            "frequency_penalty": float(params.get("freq_penalty", 1.1)),
            "max_tokens": int(params.get("max_len", 512)),
        }


        print("\n===== PROMPT SENT TO MODEL =====\n")
        print(full_prompt)
        print("\n===== END OF PROMPT =====\n")
        
        assistant_answer = call_language_model_sync(
            messages=[{"role": "user", "content": full_prompt}],
            temperature=options["temperature"],
            max_tokens=options["max_tokens"],
            model_name=model_name
        )
#        print("\n========== Debug: Full Prompt to LLM ==========")
#        print(full_prompt)
#        print("===============================================\n")

        try:
            extend_sug = generate_question_followup_suggestions(question=prompt, lang=lang) or ""
        except Exception as e:
            extend_sug = f"(Failed to generate expansion suggestion: {e})"

        try:
            followup_q = generate_followup_question(user_question, extracted_info, lang) or ""
        except Exception:
            followup_q = ""
#        print("\n========== Debug: Follow-up Question ==========")
#        print(followup_q or "(None)")
#        print("===============================================\n")

#        print("========== Debug: Extended Suggestions ==========")
#        print(extend_sug or "(None)")
#        print("=================================================\n")

        parts = [assistant_answer.strip()]
        if followup_q:
            parts.append(f"\n\nSuggested helpful next question:\n{followup_q}")
        if extend_sug:
            parts.append(f"\n\nExtended suggestions \n\n(⚠️Below response is generated by a language model. Please verify before use.):\n\n{extend_sug}")

        full_answer = "\n".join(parts).strip()

#        print("\n========== Debug: Final Answer to User ==========")
#        print(full_answer)
#        print("==================================================\n")

        ui_queue.put(full_answer)   

        if memory_mode:
            lc_memory.chat_memory.add_user_message(user_question)
            lc_memory.chat_memory.add_ai_message(full_answer)


    except Exception as exc:
        ui_queue.put(f"\n\n[ERROR] {exc}")

    finally:
        ui_queue.put(None)

def start_generation_thread(
    prompt: str,
    lang: str,
    model_name: str,
    memory_mode: bool,
    doc_paths: List[str],
    df_dict: dict | None,
    ui_queue: queue.Queue,
    params: dict,
) -> None:
    threading.Thread(
        target=_generate_in_thread,
        args=(prompt, lang, model_name,  memory_mode, doc_paths,  df_dict, ui_queue, params),
        daemon=True,
    ).start()

def format_lc_history(memory_obj) -> str:
    """Format ConversationBufferMemory into readable dialogue"""
    messages = memory_obj.chat_memory.messages
    dialogue = []
    for m in messages:
        if hasattr(m, "content") and hasattr(m, "type"):
            role = "User" if m.type == "human" else "Assistant"
            dialogue.append(f"{role}: {m.content}")
    return "\n".join(dialogue)

def call_language_model_sync(
    *,                   
    messages: list[dict],
    temperature: float = 0.7,
    max_tokens: int = 400,
    model_name: str = "llama3",
    **kwargs,
) -> str:

    prompt = "\n".join(m["content"] for m in messages)

#    print("\n[LLM Prompt Preview from call_language_model_sync()]")
#    print(prompt)
#    print("---------------------------------------------------\n")

    options = {
        "temperature": temperature,
        "max_tokens": max_tokens,    
    }

    q: queue.Queue[str | None] = queue.Queue()

    call_language_model_stream(
        full_prompt=prompt,
        model_name=model_name,
        options=options,
        beginner_tip="",      
        q=q,
    )

    parts: list[str] = []
    while True:
        token = q.get()
        if token is None:
            break
        parts.append(token)
    return "".join(parts).strip()

def generate_llm_recommendations(summary: str, lang: str = "zh") -> str:

    lang_map = {"English": "en", "中文": "zh", "日本語": "ja"}
    lang_code = lang_map.get(lang, "en")  
    model_lang_hint = {"en": "Please answer in English.",
                       "zh": "请使用中文回答。",
                       "ja": "日本語で答えてください。"}

    system_msg = {
        "zh": (
            "你是 HiokiAssist，一名电能质量工程专家。\n"
            "你的任务是根据摘要内容，提出具体可行的技术建议。\n"
            "重要：你只能用中文回答，不得使用英文或日文。\n"
            "请勿评论、讨论或比较任何非 Hioki 品牌，如 Fluke、Keysight、NI、Omron 等。\n"
            "如遇无关问题，请回复：“很抱歉，我只能协助解答与 Hioki 产品和服务相关的问题。”"
        ),
        "en": (
            "You are HiokiAssist, a power quality expert.\n"
            "Your job is to generate actionable technical recommendations from a summary.\n"
            "Important: You must respond in English only. Do not use other languages.\n"
            "Do not discuss or compare with any non-Hioki brands like Fluke, Keysight, NI, or Omron.\n"
            "If the question is irrelevant, say: 'Sorry, I can only help with Hioki-related topics.'"
        ),
        "ja": (
            "あなたは HiokiAssist という電力品質エンジニアです。\n"
            "要約に基づいて、実用的な改善提案を提示してください。\n"
            "注意：必ず日本語で回答してください。他の言語は使わないでください。\n"
            "Fluke、Keysight、NI、Omron などの他社ブランドについてコメント・比較しないでください。\n"
            "無関係な質問には、「申し訳ありませんが、Hioki 製品およびサービスに関するご質問のみお手伝いできます。」と答えてください。"
        ),
    }[lang_code]

    user_msg = {
        "zh": f"{model_lang_hint['zh']}\n以下是一份电能质量报告摘要，请给出 3-5 条可操作建议，重点关注 THD、Dip、Inrush 等指标：\n\n{summary}",
        "en": f"{model_lang_hint['en']}\nHere is a power-quality report summary. Provide 3–5 actionable recommendations focusing on THD, dips, inrush, etc.:\n\n{summary}",
        "ja": f"{model_lang_hint['ja']}\n以下は電力品質レポートの要約です。THDや電圧ディップ、インラッシュなどに注目して、3〜5件の改善提案を挙げてください：\n\n{summary}",
    }[lang_code]

    rec_text = call_language_model_sync(
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg}
        ],
        temperature=0.7,
        max_tokens=400,
        model_name="HiokiAssist", 
    )

    if not rec_text.strip():
        return {
            "zh": "（未能生成建议，请检查摘要内容是否过长或过杂）",
            "en": "(No recommendations generated. Try simplifying the summary.)",
            "ja": "（推奨事項は生成されませんでした。要約を簡素化してください）",
        }[lang]

    return rec_text.strip()

def generate_question_followup_suggestions(question: str, lang: str = "zh") -> str:

    lang_alias_to_key = {
        "English": "en", "英文": "en", "en": "en",
        "中文": "zh", "zh": "zh", "汉语": "zh",
        "日本語": "ja", "日语": "ja", "ja": "ja"
    }
    lang_key = lang_alias_to_key.get(lang, "en")

    system_prompt = {
        "zh": (
            "你是 HiokiAssist，一个电能质量专家，职责是根据用户问题提出相关建议。\n"
            "重要：你只能使用中文回答，不得使用英文或其他语言。\n"
            "请勿评论或比较任何非 Hioki 品牌，例如 Fluke、Keysight、NI、Omron。\n"
            "若问题与 Hioki 无关，请礼貌拒答："
            "“很抱歉，我只能协助解答与 Hioki 产品和服务相关的问题。”"
        ),
        "en": (
            "You are HiokiAssist, a power quality expert.\n"
            "You must respond in English only.\n"
            "Do not comment on or compare with any non-Hioki brands such as Fluke, Keysight, NI, or Omron.\n"
            "If the question is unrelated to Hioki, politely say: "
            "'Sorry, I can only assist with Hioki-related topics.'"
        ),
        "ja": (
            "あなたはHiokiAssistという電力品質の専門家で、ユーザーの質問に基づいて助言を行う役割です。\n"
            "注意：日本語のみで回答してください。他言語は禁止です。\n"
            "Fluke、Keysight、NI、Omronなどの他社製品に言及・比較しないでください。\n"
            "Hiokiに関係のない質問については、次のように丁寧に返答してください："
            "「申し訳ありませんが、Hioki 製品およびサービスに関するご質問のみお手伝いできます。」"
        )
    }[lang_key]

    user_prompt = {
        "zh": f"请你仅使用中文回答以下问题：\n\n用户问题：{question}\n\n请提供 3-5 条相关建议、注意事项或后续思考点。",
        "en": f"Please answer in English only.\n\nUser question: {question}\n\nProvide 3–5 related suggestions, cautions, or possible follow-up points.",
        "ja": f"以下はユーザーの質問です。日本語でのみ回答してください：\n\n{question}\n\n3〜5件の関連アドバイスや注意事項、フォローアップの視点を挙げてください。"
    }[lang_key]

    #print("\n========== Debug: Suggestion Prompts ==========")
    #print(f"[System Prompt]:\n{system_prompt}")
    #print(f"[User Prompt]:\n{user_prompt}")
    #print("================================================\n")

    return call_language_model_sync(
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        temperature=0.7,
        max_tokens=400,
        model_name="HiokiAssist",  
    ).strip()

def append_summary_to_docx(doc_path: str, summary_text: str) -> str:
    doc = Document(doc_path)
    doc.add_page_break()

    p = doc.add_paragraph()
    run = p.add_run("Auto-generated Summary")
    run.bold = True
    run.font.size = None 

    doc.add_paragraph(summary_text)

    new_path = doc_path.replace(".docx", "_with_summary.docx")
    doc.save(new_path)
    return new_path